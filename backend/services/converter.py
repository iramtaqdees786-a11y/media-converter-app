"""
Service for converting files between different formats.
Supports video, audio, images, documents, and spreadsheets.
"""

import asyncio
import subprocess
import os
from pathlib import Path
from typing import Optional, Tuple
from dataclasses import dataclass
from io import BytesIO
import logging

# Configure logger
logger = logging.getLogger('converter')
logger.setLevel(logging.DEBUG)
handler = logging.StreamHandler()
handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
logger.addHandler(handler)

# Image processing
from PIL import Image

# Document processing
from docx import Document
from docx.shared import Inches
from pypdf import PdfReader, PdfWriter
import pdfplumber

# Spreadsheet processing
import pandas as pd
from openpyxl import Workbook, load_workbook

from backend.config import CONVERTED_DIR, FFMPEG_PATH
from backend.utils.helpers import (
    generate_unique_filename, 
    get_file_extension, 
    get_file_category,
    validate_conversion,
    run_command_async
)


@dataclass
class ConversionResult:
    """Result of a conversion operation."""
    success: bool
    message: str
    original_file: Optional[str] = None
    converted_file: Optional[str] = None
    original_size: Optional[int] = None
    converted_size: Optional[int] = None


class VideoAudioConverter:
    """Handle video and audio conversions using FFmpeg."""
    
    @staticmethod
    async def convert(
        input_path: Path,
        output_format: str,
        quality: str = "medium"
    ) -> ConversionResult:
        """
        Convert video or audio file using FFmpeg with universal fallback.
        
        Args:
            input_path: Path to input file
            output_format: Target format (mp4, mp3, mov, webm, avi, mkv, etc.)
            quality: Quality preset - low, medium, high
        
        Returns:
            ConversionResult with conversion status
        """
        if not input_path.exists():
            return ConversionResult(
                success=False,
                message=f"Input file not found: {input_path}"
            )
        
        output_filename = generate_unique_filename(input_path.name, output_format)
        output_path = CONVERTED_DIR / output_filename
        
        input_ext = get_file_extension(input_path.name)
        input_category = get_file_category(input_ext)
        output_category = get_file_category(output_format)
        
        # Quality presets - OPTIMIZED FOR SPEED
        quality_presets = {
            "low": {"video_bitrate": "1M", "audio_bitrate": "128k", "crf": "28", "preset": "veryfast"},
            "medium": {"video_bitrate": "2M", "audio_bitrate": "192k", "crf": "23", "preset": "fast"},
            "high": {"video_bitrate": "5M", "audio_bitrate": "320k", "crf": "18", "preset": "fast"}
        }
        preset = quality_presets.get(quality, quality_presets["medium"])
        
        # Build FFmpeg command - OPTIMIZED FOR SPEED
        cmd = [FFMPEG_PATH, '-i', str(input_path), '-y', '-loglevel', 'error', '-threads', '0']
        
        # AUDIO CONVERSIONS (from any source)
        if output_category == "audio":
            cmd.extend(['-vn'])  # No video
            
            if output_format == "mp3":
                cmd.extend(['-acodec', 'libmp3lame', '-b:a', preset['audio_bitrate'], '-ar', '44100', '-q:a', '2'])
            elif output_format == "wav":
                cmd.extend(['-acodec', 'pcm_s16le', '-ar', '44100'])
            elif output_format == "aac":
                cmd.extend(['-acodec', 'aac', '-b:a', preset['audio_bitrate'], '-ar', '44100'])
            elif output_format == "ogg":
                cmd.extend(['-acodec', 'libvorbis', '-b:a', preset['audio_bitrate'], '-q:a', '4'])
            elif output_format == "flac":
                cmd.extend(['-acodec', 'flac', '-ar', '44100', '-compression_level', '5'])
            else:
                # Universal fallback for unknown audio formats
                cmd.extend(['-acodec', 'libmp3lame', '-b:a', preset['audio_bitrate'], '-q:a', '2'])
        
        # VIDEO CONVERSIONS
        elif output_category == "video":
            if output_format == "mp4":
                cmd.extend([
                    '-c:v', 'libx264',
                    '-preset', preset['preset'],
                    '-crf', preset['crf'],
                    '-c:a', 'aac',
                    '-b:a', preset['audio_bitrate'],
                    '-movflags', '+faststart',
                    '-tune', 'fastdecode'
                ])
            
            elif output_format == "mov":
                cmd.extend([
                    '-c:v', 'libx264',
                    '-preset', preset['preset'],
                    '-crf', preset['crf'],
                    '-c:a', 'aac',
                    '-b:a', preset['audio_bitrate'],
                    '-movflags', '+faststart',
                    '-pix_fmt', 'yuv420p',
                    '-tune', 'fastdecode'
                ])
            
            elif output_format == "webm":
                # Fast WebM encoding
                cmd.extend([
                    '-c:v', 'libvpx-vp9',
                    '-crf', '33',
                    '-b:v', '0',
                    '-c:a', 'libopus',
                    '-b:a', '128k',
                    '-deadline', 'realtime',
                    '-cpu-used', '4',
                    '-row-mt', '1'
                ])
            
            elif output_format == "mkv":
                cmd.extend([
                    '-c:v', 'libx264',
                    '-preset', preset['preset'],
                    '-crf', preset['crf'],
                    '-c:a', 'aac',
                    '-b:a', preset['audio_bitrate']
                ])
            
            elif output_format == "avi":
                cmd.extend([
                    '-c:v', 'mpeg4',
                    '-b:v', preset['video_bitrate'],
                    '-c:a', 'libmp3lame',
                    '-b:a', preset['audio_bitrate'],
                    '-q:v', '3'
                ])
            
            else:
                # Universal fallback for ANY video format - FAST
                cmd.extend([
                    '-c:v', 'libx264',
                    '-preset', 'veryfast',
                    '-crf', '23',
                    '-c:a', 'aac',
                    '-b:a', '192k',
                    '-pix_fmt', 'yuv420p'
                ])
        
        cmd.append(str(output_path))
        
        try:
            returncode, stdout, stderr = await run_command_async(cmd)
            
            if returncode != 0:
                # Try a second attempt with FASTEST possible settings
                print(f"First attempt failed: {stderr[:200]}")
                
                # FALLBACK COMMAND - MAXIMUM SPEED
                fallback_cmd = [
                    FFMPEG_PATH, '-i', str(input_path), '-y',
                    '-loglevel', 'error', '-threads', '0'
                ]
                
                if output_category == "audio":
                    fallback_cmd.extend([
                        '-vn', '-acodec', 'libmp3lame',
                        '-b:a', '192k', '-ar', '44100', '-q:a', '2'
                    ])
                else:
                    fallback_cmd.extend([
                        '-c:v', 'libx264', '-preset', 'ultrafast',
                        '-crf', '23', '-c:a', 'aac',
                        '-b:a', '192k', '-pix_fmt', 'yuv420p',
                        '-tune', 'zerolatency'
                    ])
                
                fallback_cmd.append(str(output_path))
                
                returncode2, stdout2, stderr2 = await run_command_async(fallback_cmd)
                
                if returncode2 != 0:
                    error_lower = stderr2.lower()
                    if 'codec not found' in error_lower or 'encoder' in error_lower:
                        friendly_msg = "Server is updating. Please try again in a moment!"
                    elif 'no such file' in error_lower or 'invalid' in error_lower:
                        friendly_msg = "The file appears to be corrupted. Please try uploading again."
                    else:
                        friendly_msg = "Conversion format not available right now. Try a different format!"
                    
                    return ConversionResult(success=False, message=friendly_msg)
            
            if not output_path.exists():
                return ConversionResult(
                    success=False,
                    message="Server is processing too many requests. Please try again!"
                )
            
            return ConversionResult(
                success=True,
                message=f"✅ Successfully converted to {output_format.upper()}! Your file is ready.",
                original_file=input_path.name,
                converted_file=output_path.name,
                original_size=input_path.stat().st_size,
                converted_size=output_path.stat().st_size
            )
            
        except FileNotFoundError:
            return ConversionResult(
                success=False,
                message="Video processing service is temporarily offline. Please try again!"
            )
        except asyncio.TimeoutError:
            return ConversionResult(
                success=False,
                message="The file is taking longer than expected. Try a smaller file!"
            )
        except Exception as e:
            # Log full traceback for debugging
            logger.exception('Unexpected error during conversion')
            
            # Provide more specific user-friendly messages based on exception type
            if isinstance(e, FileNotFoundError):
                friendly_msg = "Server configuration error: FFmpeg binary not found."
            elif isinstance(e, PermissionError):
                friendly_msg = "Server cannot write to the output directory."
            else:
                friendly_msg = "Something went wrong! We're working on it. Try again shortly."
            return ConversionResult(
                success=False,
                message=friendly_msg,
            )


class ImageConverter:
    """Handle image format conversions using Pillow."""
    
    @staticmethod
    async def convert(
        input_path: Path,
        output_format: str,
        quality: int = 95,
        resize: Optional[Tuple[int, int]] = None
    ) -> ConversionResult:
        """
        Convert image to different format.
        
        Args:
            input_path: Path to input image
            output_format: Target format (jpg, jpeg, png, webp, etc.)
            quality: JPEG/WebP quality (1-100)
            resize: Optional (width, height) tuple for resizing
        
        Returns:
            ConversionResult with conversion status
        """
        if not input_path.exists():
            return ConversionResult(
                success=False,
                message=f"Input file not found: {input_path}"
            )
        
        # Normalize format (jpeg -> jpg for PIL)
        output_format_normalized = output_format.lower()
        if output_format_normalized == 'jpeg':
            output_format_normalized = 'jpg'
            save_format = 'JPEG'
        elif output_format_normalized == 'jpg':
            save_format = 'JPEG'
        else:
            save_format = output_format_normalized.upper()
        
        output_filename = generate_unique_filename(input_path.name, output_format)
        output_path = CONVERTED_DIR / output_filename
        
        try:
            loop = asyncio.get_event_loop()
            
            def do_conversion():
                try:
                    with Image.open(input_path) as img:
                        # Convert image mode if necessary
                        if save_format in ['JPEG', 'JPG']:
                            # JPEG doesn't support transparency
                            if img.mode in ['RGBA', 'LA', 'P']:
                                # Create white background
                                background = Image.new('RGB', img.size, (255, 255, 255))
                                if img.mode == 'P':
                                    img = img.convert('RGBA')
                                if img.mode in ['RGBA', 'LA']:
                                    background.paste(img, mask=img.split()[-1])
                                    img = background
                                else:
                                    img = img.convert('RGB')
                            elif img.mode != 'RGB':
                                img = img.convert('RGB')
                        
                        elif save_format == 'PNG':
                            # PNG supports transparency
                            if img.mode == 'P':
                                img = img.convert('RGBA')
                            elif img.mode not in ['RGB', 'RGBA']:
                                img = img.convert('RGBA')
                        
                        elif save_format == 'WEBP':
                            # WebP supports both RGB and RGBA
                            if img.mode == 'P':
                                img = img.convert('RGBA')
                        
                        elif save_format == 'BMP':
                            # BMP doesn't support transparency
                            if img.mode in ['RGBA', 'LA', 'P']:
                                img = img.convert('RGB')
                        
                        elif save_format == 'GIF':
                            # Handle GIF specially
                            if hasattr(img, 'n_frames') and img.n_frames > 1:
                                # Animated GIF
                                img.save(output_path, save_all=True, optimize=True, format='GIF')
                                return
                            else:
                                # Convert to palette mode
                                if img.mode != 'P':
                                    img = img.convert('P', palette=Image.ADAPTIVE)
                        
                        # Resize if requested
                        if resize:
                            img = img.resize(resize, Image.Resampling.LANCZOS)
                        
                        # Save with format-specific options - OPTIMIZED FOR SPEED
                        save_kwargs = {'format': save_format}
                        
                        if save_format in ['JPEG', 'JPG']:
                            save_kwargs.update({
                                'quality': quality,
                                'optimize': False,  # Faster without optimize
                                'progressive': False  # Faster
                            })
                        elif save_format == 'WEBP':
                            save_kwargs.update({
                                'quality': quality,
                                'method': 4  # Faster than 6
                            })
                        elif save_format == 'PNG':
                            save_kwargs.update({
                                'optimize': False,  # Much faster
                                'compress_level': 6  # Balanced (9 is slow)
                            })
                        
                        img.save(output_path, **save_kwargs)
                
                except Exception as e:
                    raise Exception(f"Image processing error: {str(e)}")
            
            await loop.run_in_executor(None, do_conversion)
            
            if not output_path.exists():
                return ConversionResult(
                    success=False,
                    message="Image conversion failed. File not created."
                )
            
            return ConversionResult(
                success=True,
                message=f"✅ Successfully converted to {output_format.upper()}!",
                original_file=input_path.name,
                converted_file=output_path.name,
                original_size=input_path.stat().st_size,
                converted_size=output_path.stat().st_size
            )
            
        except Exception as e:
            print(f"Image conversion error: {str(e)}")
            return ConversionResult(
                success=False,
                message=f"Image conversion failed. Please try a different format or file."
            )


class DocumentConverter:
    """Handle document format conversions."""
    
    @staticmethod
    async def convert(
        input_path: Path,
        output_format: str
    ) -> ConversionResult:
        """
        Convert documents between PDF, DOCX, TXT, and XLSX formats.
        
        Args:
            input_path: Path to input document
            output_format: Target format (pdf, docx, txt, xlsx, csv)
        
        Returns:
            ConversionResult with conversion status
        """
        if not input_path.exists():
            return ConversionResult(
                success=False,
                message=f"Input file not found: {input_path}"
            )
        
        input_ext = get_file_extension(input_path.name)
        output_filename = generate_unique_filename(input_path.name, output_format)
        output_path = CONVERTED_DIR / output_filename
        
        loop = asyncio.get_event_loop()
        
        try:
            if input_ext == 'txt':
                if output_format == 'docx':
                    result = await loop.run_in_executor(
                        None, 
                        DocumentConverter._txt_to_docx, 
                        input_path, 
                        output_path
                    )
                elif output_format == 'pdf':
                    # For TXT to PDF, we go through DOCX first
                    temp_docx = CONVERTED_DIR / f"temp_{output_filename}.docx"
                    await loop.run_in_executor(
                        None, 
                        DocumentConverter._txt_to_docx, 
                        input_path, 
                        temp_docx
                    )
                    # Note: Full PDF conversion would require additional libraries like reportlab
                    # For simplicity, we'll create a basic text file labeled as conversion note
                    return ConversionResult(
                        success=False,
                        message="TXT to PDF conversion requires additional setup. Please convert to DOCX first."
                    )
            
            elif input_ext == 'docx':
                if output_format == 'txt':
                    await loop.run_in_executor(
                        None,
                        DocumentConverter._docx_to_txt,
                        input_path,
                        output_path
                    )
                elif output_format == 'pdf':
                    return ConversionResult(
                        success=False,
                        message="DOCX to PDF conversion requires LibreOffice or similar. Please use an online converter."
                    )
            
            elif input_ext == 'pdf':
                if output_format == 'txt':
                    await loop.run_in_executor(
                        None,
                        DocumentConverter._pdf_to_txt,
                        input_path,
                        output_path
                    )
                elif output_format == 'docx':
                    # Basic PDF text extraction to DOCX
                    await loop.run_in_executor(
                        None,
                        DocumentConverter._pdf_to_docx,
                        input_path,
                        output_path
                    )
                elif output_format in ['xlsx', 'xls', 'csv']:
                    # PDF to Excel/CSV (Table extraction)
                    await loop.run_in_executor(
                        None,
                        DocumentConverter._pdf_to_xlsx,
                        input_path,
                        output_path,
                        output_format
                    )
            
            if output_path.exists():
                return ConversionResult(
                    success=True,
                    message=f"Successfully converted to {output_format}",
                    original_file=input_path.name,
                    converted_file=output_path.name,
                    original_size=input_path.stat().st_size,
                    converted_size=output_path.stat().st_size
                )
            else:
                return ConversionResult(
                    success=False,
                    message="Conversion failed - output file not created"
                )
                
        except Exception as e:
            return ConversionResult(
                success=False,
                message=f"Document conversion error: {str(e)}"
            )
    
    @staticmethod
    def _txt_to_docx(input_path: Path, output_path: Path):
        """Convert TXT to DOCX."""
        doc = Document()
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                doc.add_paragraph(line.strip())
        doc.save(output_path)
    
    @staticmethod
    def _docx_to_txt(input_path: Path, output_path: Path):
        """Convert DOCX to TXT."""
        doc = Document(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')
    
    @staticmethod
    def _pdf_to_txt(input_path: Path, output_path: Path):
        """Extract text from PDF."""
        reader = PdfReader(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    f.write(text + '\n')
    
    @staticmethod
    def _pdf_to_docx(input_path: Path, output_path: Path):
        """Convert PDF to DOCX (basic text extraction)."""
        reader = PdfReader(input_path)
        doc = Document()
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
            doc.add_page_break()
        
        doc.save(output_path)
    
    @staticmethod
    def _pdf_to_xlsx(input_path: Path, output_path: Path, output_format: str):
        """Convert PDF to Excel/CSV by extracting tables."""
        tables = []
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_tables()
                for table in extracted:
                    # Clean up table data: remove None values
                    cleaned_table = [[cell if cell is not None else "" for cell in row] for row in table]
                    if cleaned_table:
                        df = pd.DataFrame(cleaned_table)
                        # Assume first row is header if it looks like one, otherwise just data
                        # For simplicity, we just dump the data
                        tables.append(df)
        
        if not tables:
            # Fallback: if no tables found, try to extract text and put in one cell per page?
            # Or just raise error/empty file
            # Let's create an empty dataframe with a message
            df = pd.DataFrame(["No tables found in PDF"], columns=["Message"])
            tables.append(df)

        # Concatenate all tables or put them in separate sheets?
        # For CSV, we must concatenate. For Excel, we could do sheets.
        # Let's simple concatenate for version 1
        final_df = pd.concat(tables, ignore_index=True) if tables else pd.DataFrame()
        
        if output_format == 'csv':
            final_df.to_csv(output_path, index=False, header=False)
        else:
            final_df.to_excel(output_path, index=False, header=False)



class SpreadsheetConverter:
    """Handle spreadsheet format conversions."""
    
    @staticmethod
    async def convert(
        input_path: Path,
        output_format: str
    ) -> ConversionResult:
        """
        Convert spreadsheets between XLSX, XLS, and CSV formats.
        
        Args:
            input_path: Path to input spreadsheet
            output_format: Target format (xlsx, xls, csv)
        
        Returns:
            ConversionResult with conversion status
        """
        if not input_path.exists():
            return ConversionResult(
                success=False,
                message=f"Input file not found: {input_path}"
            )
        
        input_ext = get_file_extension(input_path.name)
        output_filename = generate_unique_filename(input_path.name, output_format)
        output_path = CONVERTED_DIR / output_filename
        
        loop = asyncio.get_event_loop()
        
        try:
            def do_conversion():
                # Read input file
                if input_ext == 'csv':
                    df = pd.read_csv(input_path)
                elif input_ext in ['xlsx', 'xls']:
                    df = pd.read_excel(input_path, engine='openpyxl' if input_ext == 'xlsx' else 'xlrd')
                else:
                    raise ValueError(f"Unsupported input format: {input_ext}")
                
                # Write output file
                if output_format == 'csv':
                    df.to_csv(output_path, index=False)
                elif output_format == 'xlsx':
                    df.to_excel(output_path, index=False, engine='openpyxl')
                elif output_format == 'xls':
                    # Note: xlwt only supports xls, but pandas defaults to openpyxl for xlsx
                    # For xls output, we save as xlsx with .xls extension (most apps can read it)
                    df.to_excel(output_path, index=False, engine='openpyxl')
            
            await loop.run_in_executor(None, do_conversion)
            
            return ConversionResult(
                success=True,
                message=f"Successfully converted to {output_format}",
                original_file=input_path.name,
                converted_file=output_path.name,
                original_size=input_path.stat().st_size,
                converted_size=output_path.stat().st_size
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                message=f"Spreadsheet conversion error: {str(e)}"
            )


async def convert_file(
    input_path: Path,
    output_format: str,
    **kwargs
) -> ConversionResult:
    """
    Main conversion function that routes to appropriate converter.
    
    Args:
        input_path: Path to input file
        output_format: Target format
        **kwargs: Additional options for specific converters
    
    Returns:
        ConversionResult with conversion status
    """
    input_ext = get_file_extension(input_path.name)
    
    # Validate conversion
    is_valid, error_msg = validate_conversion(input_ext, output_format)
    if not is_valid:
        return ConversionResult(success=False, message=error_msg)
    
    input_category = get_file_category(input_ext)
    output_category = get_file_category(output_format)
    
    # Route to appropriate converter
    if input_category in ['video', 'audio'] or output_category in ['video', 'audio']:
        return await VideoAudioConverter.convert(
            input_path, 
            output_format,
            quality=kwargs.get('quality', 'medium')
        )
    
    elif input_category == 'image':
        return await ImageConverter.convert(
            input_path,
            output_format,
            quality=kwargs.get('quality', 85),
            resize=kwargs.get('resize')
        )
    
    elif input_category == 'document':
        return await DocumentConverter.convert(input_path, output_format)
    
    elif input_category == 'spreadsheet':
        return await SpreadsheetConverter.convert(input_path, output_format)
    
    else:
        return ConversionResult(
            success=False,
            message=f"No converter available for {input_category} files"
        )
